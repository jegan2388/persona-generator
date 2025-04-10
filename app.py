from flask import Flask, request, jsonify, send_file
import requests
from bs4 import BeautifulSoup
import os
from dotenv import load_dotenv
from pathlib import Path
import logging
import sys
import json
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from flask_cors import CORS

# Configure logging to output to stdout
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[logging.StreamHandler(sys.stdout)]
)
logger = logging.getLogger(__name__)

# Log startup information
logger.info("Starting application...")
logger.info(f"Python version: {sys.version}")
logger.info(f"Current working directory: {os.getcwd()}")

try:
    # Load .env file
    load_dotenv(dotenv_path=Path(".env"))
    logger.info("Loaded .env file")

    # Get and set the OpenAI API key in environment
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise ValueError("No OpenAI API key found in environment variables")
    os.environ["OPENAI_API_KEY"] = api_key
    logger.info("OpenAI API key found")

    # Import OpenAI and configure HTTP client
    import httpx
    from openai import OpenAI
    
    # Create a custom HTTP client without proxy settings
    http_client = httpx.Client(
        base_url="https://api.openai.com/v1",
        follow_redirects=True
    )
    
    # Initialize OpenAI with custom client
    client = OpenAI(
        http_client=http_client
    )
    logger.info("OpenAI client initialized")

    app = Flask(__name__)
    logger.info("Flask app created")

    # Configure CORS based on environment
    if os.getenv('FLASK_ENV') == 'production':
        allowed_origins = ['*']  # Allow all origins in production
        logger.info("Production CORS: allowing all origins")
    else:
        allowed_origins = [
            "http://localhost:8000",
            "http://127.0.0.1:8000",
            "http://localhost:5000",
            "http://127.0.0.1:5000",
            "http://[::]:8000",
            "http://[::1]:8000"
        ]
        logger.info("Development CORS origins configured")

    CORS(app, resources={
        r"/*": {
            "origins": allowed_origins,
            "methods": ["GET", "POST", "OPTIONS"],
            "allow_headers": ["Content-Type"]
        }
    })
    logger.info("CORS configured")

except Exception as e:
    logger.error(f"Error during initialization: {str(e)}")
    raise

def scrape_website(url):
    try:
        response = requests.get(url, timeout=10)
        soup = BeautifulSoup(response.text, "html.parser")
        return soup.get_text(separator=" ", strip=True)
    except Exception as e:
        return f"Error scraping site: {e}"

def generate_persona(content, job_title, industry="Technology"):
    prompt = f"""
You are a B2B persona strategist with 15+ years of experience working with GTM, sales, and marketing teams to develop deep, actionable customer personas. You specialize in turning limited inputs into focused, insightful outputs that help teams align messaging, campaigns, and product strategy.

Your job is to take a website/product description and a target job title, and generate a clear, well-structured customer persona tailored to that role.

---

🧭 Guidelines:

- Do not write in paragraph format. 
- Use **bullet points only**, with clear headers for each section.
- Ensure each section is **distinct** — do not repeat ideas across sections (e.g., don't let "Goals" and "Pain Points" overlap).
- Focus on **practical, job-specific details**, not generic fluff.
- Use the scraped content as product context to make the persona relevant.
- Avoid marketing buzzwords (e.g., "synergy", "supercharge", "delightful").
- Write in a **confident, professional tone** that would feel natural in a strategy or planning document.
- Tailor everything to the job title and industry — assume this is for a real B2B team.

---

🎯 Format your output as follows (exact headers, in this order):

**Persona Name:**  
(A short fictional name for this persona)

**Job Titles:**  
- (Other job titles this persona might also hold)

**Background:**  
- (Their experience, company type, industry exposure, education if relevant)

**Responsibilities:**  
- (Key day-to-day activities they manage)

**Pain Points:**  
- (Frustrations, blockers, risks they face regularly)

**Goals:**  
- (What success looks like in their role)

**Objections to Our Tool:**  
- (Reasons they might hesitate to adopt our product)

**How Our Tool Helps:**  
- (How the product specifically solves for their pain points and supports their goals)

---

Here is the product context:

{content[:3000]}

The target persona is a {job_title} in the {industry} industry.
"""

    response = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7,
    )

    return response.choices[0].message.content

def analyze_persona_traits(persona_text):
    prompt = f"""
Analyze the following persona description and rate their key behavioral traits on a scale of 1-10.
Return the results in JSON format.

Traits to analyze:
1. Strategic Thinking: How well they plan and think long-term
2. Tech Savviness: Their comfort and proficiency with technology
3. Risk Aversion: How cautious they are with new tools/approaches
4. Speed of Decision Making: How quickly they make decisions

IMPORTANT: Return ONLY a valid JSON object with these exact keys:
{{
    "strategic_thinking": number,
    "tech_savviness": number,
    "risk_aversion": number,
    "decision_speed": number,
    "trait_descriptions": {{
        "strategic_thinking": "brief description",
        "tech_savviness": "brief description",
        "risk_aversion": "brief description",
        "decision_speed": "brief description"
    }}
}}

Do not include any other text or explanation. Only return the JSON object.

Persona description:
{persona_text}
"""

    response = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7
    )

    try:
        # Extract the JSON string from the response
        json_str = response.choices[0].message.content.strip()
        # Parse the JSON string
        return json_str
    except Exception as e:
        logger.error(f"Error parsing traits JSON: {str(e)}")
        # Return a default response if parsing fails
        return json.dumps({
            "strategic_thinking": 5,
            "tech_savviness": 5,
            "risk_aversion": 5,
            "decision_speed": 5,
            "trait_descriptions": {
                "strategic_thinking": "Moderate strategic thinking ability",
                "tech_savviness": "Average tech proficiency",
                "risk_aversion": "Moderate risk tolerance",
                "decision_speed": "Average decision-making speed"
            }
        })

@app.route('/generate-persona', methods=['POST'])
def generate_persona_route():
    try:
        data = request.get_json()
        website_url = data.get('website')
        job_title = data.get('jobTitle')
        industry = data.get('industry', 'Technology')

        if not website_url or not job_title:
            return jsonify({"error": "Missing required fields"}), 400

        # Scrape website content
        content = scrape_website(website_url)
        
        # Generate persona
        persona = generate_persona(content, job_title, industry)
        
        # Analyze traits
        traits = analyze_persona_traits(persona)
        
        return jsonify({
            "persona": persona,
            "traits": traits
        })

    except Exception as e:
        logger.error(f"Error generating persona: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/download-docx', methods=['POST'])
def download_docx():
    try:
        data = request.get_json()
        persona_text = data.get('persona')
        traits = data.get('traits', {})
        
        if not persona_text:
            return jsonify({"error": "Missing persona text"}), 400

        # Create a new Document
        doc = Document()
        
        # Add a title
        title = doc.add_heading('Persona Profile', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add the main persona content
        doc.add_paragraph(persona_text)
        
        # Add a page break before traits
        doc.add_page_break()
        
        # Add traits section
        traits_heading = doc.add_heading('Behavioral Traits Analysis', level=1)
        traits_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add each trait with a rating bar
        for trait, value in traits.items():
            # Format trait name
            trait_name = trait.replace('_', ' ').title()
            p = doc.add_paragraph()
            p.add_run(f'{trait_name}: ').bold = True
            p.add_run(f'{value}/10')
            
            # Add a simple text-based rating bar
            bar = '█' * int(value) + '░' * (10 - int(value))
            doc.add_paragraph(bar)
        
        # Save the document to a BytesIO object
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return send_file(
            doc_io,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='persona_profile.docx'
        )

    except Exception as e:
        logger.error(f"Error creating document: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/health', methods=['GET'])
def health_check():
    return jsonify({"status": "healthy"}), 200

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 3000))
    debug = os.environ.get('FLASK_ENV') != 'production'
    logger.info(f"Starting server on port {port} with debug={debug}")
    app.run(host='0.0.0.0', port=port, debug=debug)
